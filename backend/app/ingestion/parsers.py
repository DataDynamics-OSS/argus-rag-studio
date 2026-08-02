# SPDX-License-Identifier: Apache-2.0
"""문서 파서 프로바이더 추상화 — 원본 바이트에서 텍스트(Markdown)를 추출한다.

컬렉션마다 문서 형식·품질 요구가 다르므로, 파싱 방법을 ``parse_strategy`` 로 교체한다.
임베딩 프로바이더(``app.embedding.provider``)와 같은 패턴이며, parse 시점 설정이라
변경하면 재인덱싱이 필요하다(청킹·임베딩과 동일 등급).

전략(provider kind):
    - ``text``  : 평문 텍스트 추출(pypdf/python-docx/html.parser). 빠르고 의존성 가벼움.
                  표·다단·스캔본에 취약. 항상 사용 가능(기본값).
    - ``layout``: 레이아웃 인식 — ``pdfplumber`` 로 PDF 표를 Markdown 으로 직렬화하고
                  본문을 함께 추출. 표가 핵심인 문서에서 ``text`` 대비 크게 개선.
                  비PDF/미설치 시 ``text`` 로 폴백.
    - ``docai`` : 문서 AI(``docling``) — 레이아웃 탐지·표·읽기순서를 모델로 복원.
                  미설치 시 명확한 오류(임베딩 local 과 동일한 graceful degrade).
    - ``vlm``   : 비전 LLM — 페이지를 이미지로 렌더링(``pymupdf``)해 OpenAI 호환
                  비전 엔드포인트(기존 LLM 설정 재사용)로 Markdown 을 생성. 복잡 레이아웃·
                  스캔본에 최고 품질이나 느리고 비용·환각 리스크. 의존성/설정 없으면 오류.
    - ``rhwp``  : 한글(HWP/HWPX) 전용 — Rust 파서 ``rhwp``(MIT)의 PyO3 바인딩(``rhwp_py``)
                  으로 표(셀 병합 포함)를 Markdown 으로 보존해 추출. 순수 Python 로더(``text``)
                  보다 표·레이아웃 정확도가 높다. HWP/HWPX 외 포맷은 ``text`` 로 폴백.
                  바인딩 미빌드 시 오류(backend/native/rhwp_py 참조).

모든 파서는 ``parse(filename, data) -> str`` 단일 메서드를 구현한다.
"""

import base64
import logging

from app.ingestion.loaders import extract_text

logger = logging.getLogger(__name__)

PARSE_STRATEGIES = ("auto", "text", "layout", "docai", "vlm", "rhwp")


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    dot = name.rfind(".")
    return name[dot:] if dot >= 0 else ""


# ── text (기본) ──────────────────────────────────────────────────────────────
def _parse_text(filename: str, data: bytes, markers: bool = False) -> str:
    """현재 기본 동작 — 평문 텍스트 추출(레이아웃/표/OCR 미적용).

    ``markers=True`` 이고 PDF 면 이미지 위치 마커(``[[IMG:p…:i…]]``)를 본문에 삽입한다
    (이미지 분류 ↔ 청크 연결용). PPTX 는 슬라이드 관계로 매핑해 ``## 슬라이드 N`` 섹션에,
    DOCX 는 문단 관계로 매핑해 이미지 문단 뒤에 마커를 넣는다(비 PDF 정밀 연결 —
    design/image-content-indexing.md §6 Step 3). 삽입 실패 시 평문으로 폴백한다."""
    if markers and _ext(filename) == ".pdf":
        try:
            return _pdf_text_with_markers(data)
        except Exception:  # noqa: BLE001 — 마커 삽입 실패 시 평문으로 폴백
            logger.warning("text 파서 이미지 마커 삽입 실패 — 평문으로 폴백", exc_info=True)
    if markers and _ext(filename) == ".pptx":
        try:
            return _pptx_text_with_markers(filename, data)
        except Exception:  # noqa: BLE001 — 마커 삽입 실패 시 평문으로 폴백
            logger.warning("pptx 이미지 마커 삽입 실패 — 평문으로 폴백", exc_info=True)
    if markers and _ext(filename) == ".docx":
        try:
            return _docx_text_with_markers(filename, data)
        except Exception:  # noqa: BLE001 — 마커 삽입 실패 시 평문으로 폴백
            logger.warning("docx 이미지 마커 삽입 실패 — 평문으로 폴백", exc_info=True)
    if markers and _ext(filename) in (".html", ".htm"):
        try:
            return _html_text_with_markers(filename, data)
        except Exception:  # noqa: BLE001 — 마커 삽입 실패 시 평문으로 폴백
            logger.warning("html 이미지 마커 삽입 실패 — 평문으로 폴백", exc_info=True)
    if markers and _ext(filename) == ".hwpx":
        try:
            return _hwpx_text_with_markers(filename, data)
        except Exception:  # noqa: BLE001 — 마커 삽입 실패 시 평문으로 폴백
            logger.warning("hwpx 이미지 마커 삽입 실패 — 평문으로 폴백", exc_info=True)
    return extract_text(filename, data)


def _pdf_text_with_markers(data: bytes) -> str:
    """text 파서 PDF 경로(pypdf) + 페이지별 이미지 위치 마커 삽입."""
    import io

    from pypdf import PdfReader

    from app.ingestion.image_classifier import insert_page_markers, pdf_image_layout

    layout = pdf_image_layout(data)
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for pno, page in enumerate(reader.pages, start=1):
        t = (page.extract_text() or "").strip()
        if layout.get(pno):
            t = insert_page_markers(t, layout[pno])
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts).strip()


def _pptx_text_with_markers(filename: str, data: bytes) -> str:
    """pptx 본문(``## 슬라이드 N`` 섹션) + 슬라이드별 이미지 마커 삽입."""
    from app.ingestion.image_classifier import append_slide_markers, media_image_layout

    text = extract_text(filename, data)
    layout = media_image_layout(filename, data)
    if not layout:
        return text
    return append_slide_markers(text, layout)


def _docx_text_with_markers(filename: str, data: bytes) -> str:
    """docx 본문 + 이미지 문단 뒤 마커 삽입.

    로더(``_load_docx``)와 동일한 문단 축(python-docx ``Document.paragraphs``)을 직접
    순회하며 텍스트를 재구성한다 — 빈 문단(이미지 전용 문단 포함)은 로더처럼 생략하되,
    그 문단의 이미지 마커는 해당 위치에 남는다. 표 안 이미지는 매핑에 없어 마커 없이
    부록 폴백(캡션 주입 시 말미 섹션)으로 처리된다.
    """
    import io

    from docx import Document

    from app.ingestion.image_classifier import media_image_layout

    layout = media_image_layout(filename, data)  # {문단 번호: [{index, page}]}
    doc = Document(io.BytesIO(data))
    if not layout:
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    from app.ingestion.image_classifier import _MARKER_FMT

    lines: list[str] = []
    for pno, p in enumerate(doc.paragraphs, start=1):
        if p.text.strip():
            lines.append(p.text)
        for it in layout.get(pno, ()):
            lines.append(_MARKER_FMT.format(page=it["page"], index=it["index"]))
    return "\n".join(lines)


def _html_text_with_markers(filename: str, data: bytes) -> str:
    """html 본문 + data URI ``<img>`` 위치에 마커 삽입.

    로더의 텍스트 추출기(``_HTMLTextExtractor``)를 확장해, 추출기(``_extract_html``)와
    동일한 **data URI img 등장 순번** 축으로 해당 지점에 마커를 끼운다. 외부 URL img 는
    추출 대상이 아니므로 세지 않는다(air-gap — 순번 정합 유지).
    """
    from app.ingestion.image_classifier import _MARKER_FMT, media_image_layout
    from app.ingestion.loaders import _HTMLTextExtractor, _load_text

    layout = media_image_layout(filename, data)  # {등장 순번: [{index, page}]}
    if not layout:
        return extract_text(filename, data)

    class _WithMarkers(_HTMLTextExtractor):
        def __init__(self) -> None:
            super().__init__()
            self._img_seen = 0

        def handle_starttag(self, tag, attrs):
            super().handle_starttag(tag, attrs)
            if tag != "img":
                return
            src = (next((v for k, v in attrs if k == "src"), "") or "").strip().lower()
            # 추출기 정규식과 같은 조건(data:image + base64)만 순번에 포함.
            if not src.startswith("data:image/") or ";base64," not in src:
                return
            self._img_seen += 1
            for it in layout.get(self._img_seen, ()):
                self._parts.append(_MARKER_FMT.format(page=it["page"], index=it["index"]))

    parser = _WithMarkers()
    parser.feed(_load_text(data))
    return parser.text()


def _hwpx_text_with_markers(filename: str, data: bytes) -> str:
    """hwpx 본문(자체 OWPML 로더) + 그림(binaryItemIDRef pic) 자리에 마커 삽입.

    등장 순번 축은 추출기(``_hwpx_pic_media_map``)와 동일. text 전략 전용 —
    rhwp 전략은 외부 markdown 산출이라 마커를 끼울 수 없다(캡션은 부록 폴백).
    """
    from app.ingestion.image_classifier import _MARKER_FMT, media_image_layout
    from app.ingestion.loaders import load_hwpx_with_image_markers

    layout = media_image_layout(filename, data)  # {등장 순번: [{index, page}]}
    if not layout:
        return extract_text(filename, data)
    return load_hwpx_with_image_markers(data, layout, _MARKER_FMT)


# ── layout (pdfplumber 표 직렬화) ────────────────────────────────────────────
def _table_to_markdown(table: list[list]) -> str:
    """pdfplumber 표(행 리스트)를 GitHub Markdown 표로 직렬화한다."""
    rows = [[(c if c is not None else "").strip().replace("\n", " ") for c in row] for row in table]
    rows = [r for r in rows if any(cell for cell in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


def _parse_layout(filename: str, data: bytes, markers: bool = False) -> str:
    """PDF 는 pdfplumber 로 페이지별 본문 + 표(Markdown)를 추출. 그 외/미설치는 text 폴백.

    ``markers=True`` 면 페이지별 이미지 위치 마커를 본문에 삽입한다(이미지 분류 연결용)."""
    if _ext(filename) != ".pdf":
        return _parse_text(filename, data, markers=markers)
    try:
        import io

        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber 미설치 — layout 파서가 text 로 폴백합니다. `pip install pdfplumber`")
        return _parse_text(filename, data, markers=markers)

    layout: dict = {}
    if markers:
        from app.ingestion.image_classifier import insert_page_markers, pdf_image_layout

        layout = pdf_image_layout(data)

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if layout.get(pno):
                page_text = insert_page_markers(page_text, layout[pno])
            if page_text.strip():
                parts.append(page_text)
            for table in page.extract_tables() or []:
                md = _table_to_markdown(table)
                if md:
                    parts.append(md)
    text = "\n\n".join(parts).strip()
    if not text:
        # 텍스트 레이어가 없는 스캔본 — text 파서가 동일하게 실패하며 OCR 안내를 던진다.
        return _parse_text(filename, data)
    return text


# ── docai (Docling) ──────────────────────────────────────────────────────────
def _parse_docai(filename: str, data: bytes, markers: bool = False) -> str:  # noqa: ARG001 — docai 는 비페이지(마커 미지원)
    """Docling 문서 AI 로 Markdown 추출(레이아웃/표/읽기순서 복원). 미설치 시 오류."""
    try:
        from docling.datamodel.base_models import DocumentStream
        from docling.document_converter import DocumentConverter
    except ImportError as e:
        raise RuntimeError(
            "docai 파서에는 Docling 이 필요합니다. `pip install docling` 후 사용하세요."
        ) from e
    import io

    source = DocumentStream(name=filename or "document", stream=io.BytesIO(data))
    result = DocumentConverter().convert(source)
    text = (result.document.export_to_markdown() or "").strip()
    if not text:
        raise RuntimeError("docai 파서가 텍스트를 추출하지 못했습니다.")
    return text


# ── vlm (비전 LLM) ───────────────────────────────────────────────────────────
_VLM_PROMPT = (
    "이 문서 페이지 이미지를 빠짐없이 Markdown 으로 변환하세요. 표는 Markdown 표로, "
    "제목·목록 구조를 보존하고, 페이지 머리글/바닥글/페이지 번호는 제외하세요. "
    "설명 없이 변환 결과만 출력하세요."
)


def _parse_vlm(filename: str, data: bytes, markers: bool = False) -> str:
    """PDF 페이지를 이미지로 렌더링해 OpenAI 호환 비전 엔드포인트로 Markdown 생성.

    엔드포인트/모델은 기존 LLM 설정(``settings.llm_*``)을 재사용한다. 비전 미지원 모델이면
    엔드포인트가 오류를 반환한다. 의존성(pymupdf/httpx)·설정이 없으면 명확한 오류를 던진다.
    ``markers=True`` 면 페이지별 이미지 위치 마커를 생성 Markdown 에 삽입한다(이미지 분류 연결용)."""
    if _ext(filename) != ".pdf":
        raise RuntimeError("vlm 파서는 현재 PDF 만 지원합니다.")
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError(
            "vlm 파서에는 PyMuPDF 가 필요합니다. `pip install pymupdf` 후 사용하세요."
        ) from e
    import httpx

    from app.core.config import settings
    from app.core.http import auth_headers

    if not settings.llm_model:
        raise RuntimeError("vlm 파서에는 비전 가능한 llm.model 설정이 필요합니다.")

    base = settings.llm_server_url.rstrip("/")
    url = f"{base}/chat/completions"
    headers = {"Content-Type": "application/json"}
    headers.update(auth_headers(
        settings.llm_api_key, settings.llm_auth_header, settings.llm_auth_scheme
    ))

    layout: dict = {}
    if markers:
        from app.ingestion.image_classifier import insert_page_markers, pdf_image_layout

        layout = pdf_image_layout(data)

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc, httpx.Client(
        timeout=settings.llm_timeout, headers=headers
    ) as client:
        for pno, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=150)
            b64 = base64.b64encode(pix.tobytes("png")).decode("ascii")
            payload = {
                "model": settings.llm_model,
                "max_tokens": settings.llm_max_tokens,
                "temperature": 0,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": _VLM_PROMPT},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/png;base64,{b64}"},
                            },
                        ],
                    }
                ],
            }
            resp = client.post(url, json=payload)
            if resp.status_code != 200:
                raise RuntimeError(
                    f"비전 LLM 오류({resp.status_code}): {resp.text[:200]} @ {url}"
                )
            content = resp.json()["choices"][0]["message"]["content"]
            page_md = content.strip() if content else ""
            if layout.get(pno):
                page_md = insert_page_markers(page_md, layout[pno])
            if page_md.strip():
                parts.append(page_md)
    text = "\n\n".join(parts).strip()
    if not text:
        raise RuntimeError("vlm 파서가 텍스트를 추출하지 못했습니다.")
    return text


# ── rhwp (한글 HWP/HWPX — Rust 파서 바인딩) ──────────────────────────────────
def _parse_rhwp(filename: str, data: bytes, markers: bool = False) -> str:
    """rhwp_py(Rust 파서 PyO3 바인딩)로 HWP/HWPX 를 Markdown(표 보존)으로 추출한다.

    HWP/HWPX 외 포맷은 순수 Python ``text`` 로 폴백. 바인딩 미설치 시 명확한 오류.
    """
    if _ext(filename) not in (".hwp", ".hwpx"):
        return _parse_text(filename, data, markers=markers)
    try:
        import rhwp_py
    except ImportError as e:
        raise RuntimeError(
            "rhwp 파서에는 rhwp_py 확장이 필요합니다. backend/native/rhwp_py 를 빌드·설치하세요"
            "(`maturin build --release` 후 wheel 설치, README 참조)."
        ) from e
    # HWPX + 마커 모드: rhwp 의 이미지 토큰([[RHWP_IMAGE:n]])을 우리 위치 마커로 치환
    # (bin_data_id → content.hpf → 분류 index). 구버전 바인딩(함수 없음)은 평문 경로.
    if (
        markers
        and _ext(filename) == ".hwpx"
        and hasattr(rhwp_py, "extract_markdown_pages_with_images")
    ):
        try:
            text = _rhwp_markdown_with_markers(filename, data)
            if text:
                return text
        except Exception:  # noqa: BLE001 — 마커 삽입 실패 시 기존 경로로 폴백
            logger.warning("rhwp 이미지 마커 치환 실패 — 토큰 제거 경로로 폴백", exc_info=True)
    text = (rhwp_py.extract_markdown(data) or "").strip()
    if not text:
        raise RuntimeError("rhwp 파서가 텍스트를 추출하지 못했습니다(스캔본/암호화 가능).")
    return text


def _rhwp_markdown_with_markers(filename: str, data: bytes) -> str:
    """rhwp 페이지별 Markdown 의 이미지 토큰을 위치 마커로 치환해 합친다.

    체인: 토큰 n → 그 페이지의 bin_data_id → content.hpf(id 숫자부→미디어 멤버) →
    분류 항목의 (page, index) 마커. 같은 이미지의 반복 토큰(모든 페이지의 로고 등)은
    **첫 등장만** 마커로 치환하고 나머지는 제거한다(청크 refs 노이즈 방지 — PPTX 의
    첫 슬라이드 규칙과 동일). 매핑 안 되는 토큰도 제거(기존 strip 동작과 동일).
    """
    import re

    import rhwp_py

    from app.imaging.extractors import hwpx_binid_media_map
    from app.ingestion.image_classifier import _MARKER_FMT, media_ref_map

    pages = rhwp_py.extract_markdown_pages_with_images(data)
    binid_media = hwpx_binid_media_map(data)      # bin id → 미디어 멤버명(소문자)
    ref_of = media_ref_map(filename, data)        # 멤버명 → {page, index}
    token_re = re.compile(r"\[\[RHWP_IMAGE:(\d+)\]\]")

    seen: set[tuple[int, int]] = set()
    parts: list[str] = []
    for md, ids in pages:
        def _sub(m: "re.Match[str]") -> str:
            n = int(m.group(1))
            bid = ids[n - 1] if 1 <= n <= len(ids) else 0
            ref = ref_of.get(binid_media.get(bid, ""))
            if not ref:
                return ""
            key = (ref["page"], ref["index"])
            if key in seen:
                return ""
            seen.add(key)
            return _MARKER_FMT.format(page=ref["page"], index=ref["index"])

        page_md = token_re.sub(_sub, md).strip()
        if page_md:
            parts.append(page_md)
    return "\n\n".join(parts).strip()


_PARSERS = {
    "text": _parse_text,
    "layout": _parse_layout,
    "docai": _parse_docai,
    "vlm": _parse_vlm,
    "rhwp": _parse_rhwp,
}


def _installed(mod: str) -> bool:
    import importlib.util

    return importlib.util.find_spec(mod) is not None


def resolve_parse_strategy(filename: str) -> str:
    """파일 확장자로 *설치된* 최적 파싱 전략을 자동 선택(auto)한다.

    PDF→layout(pdfplumber 있을 때), HWP/HWPX→rhwp(rhwp_py 있을 때), 그 외/미설치→text.
    docai·vlm 은 무겁거나 외부 엔드포인트가 필요해 자동 선택하지 않는다(명시 전용)."""
    ext = _ext(filename)
    if ext == ".pdf" and _installed("pdfplumber"):
        return "layout"
    if ext in (".hwp", ".hwpx") and _installed("rhwp_py"):
        return "rhwp"
    return "text"


def parse_document(
    filename: str,
    data: bytes,
    strategy: str = "text",
    *,
    insert_image_markers: bool = False,
) -> str:
    """``strategy`` 에 따라 원본 바이트를 텍스트(Markdown)로 파싱한다(인제스천 진입점).

    ``auto`` 면 파일 유형으로 전략을 자동 선택한다(``resolve_parse_strategy``).

    ``insert_image_markers=True`` 면 PDF(text/layout/vlm)에 이미지 위치 마커(``[[IMG:p…:i…]]``)를
    삽입한다 — 인제스천이 이미지 분류 ↔ 청크 연결에 쓰고, 저장 전 본문에서 제거한다. 미리보기 등
    다른 호출자는 기본값(False)이라 마커가 노출되지 않는다.

    Raises:
        RuntimeError: 미지원 포맷/의존성 부재/추출 실패 시(워커가 잡 실패로 기록).
    """
    strat = (strategy or "text").lower()
    if strat == "auto":
        strat = resolve_parse_strategy(filename)
    fn = _PARSERS.get(strat, _parse_text)
    logger.info("문서 파싱: strategy=%s file=%s (%d bytes)", strat, filename, len(data))
    return fn(filename, data, markers=insert_image_markers)


def list_parse_strategies() -> list[dict]:
    """파싱 전략 목록 + 가용성(의존성 설치 여부)을 반환한다(UI 안내용).

    ``auto``/``text`` 는 항상 가능. ``layout`` 은 pdfplumber, ``docai`` 는 docling, ``vlm`` 은
    pymupdf 설치 여부로 판단한다(런타임 import 시도). 미설치 전략은 UI 에서 '(미설치)' 로
    표시하되 선택 자체는 허용한다(설치 후 동작)."""
    meta = [
        ("auto", "자동 (파일 유형별)", "PDF→layout, HWP/HWPX→rhwp, 그 외→text 를 파일마다 자동 선택.", True),
        ("text", "텍스트 추출 (기본)", "빠름. 표·스캔본 취약.", True),
        ("layout", "레이아웃·표 인식", "PDF 표를 Markdown 으로. pdfplumber 필요.", _installed("pdfplumber")),
        ("docai", "문서 AI (Docling)", "레이아웃·표·읽기순서 복원. docling 필요.", _installed("docling")),
        ("vlm", "비전 LLM", "페이지 이미지→Markdown. 고품질·고비용. pymupdf + 비전 모델 필요.", _installed("fitz")),
        ("rhwp", "한글 HWP/HWPX (rhwp)", "Rust 파서로 표 보존 추출. rhwp_py 빌드 필요.", _installed("rhwp_py")),
    ]
    return [
        {"strategy": s, "label": label, "description": desc, "available": avail}
        for s, label, desc, avail in meta
    ]
